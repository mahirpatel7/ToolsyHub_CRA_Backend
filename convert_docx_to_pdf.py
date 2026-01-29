# convert_docx_to_pdf.py
import sys
from docx import Document
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import os

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using reportlab"""
    try:
        # Load the Word document
        doc = Document(docx_path)
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for PDF elements
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Extract text from Word document
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine alignment
                alignment = TA_LEFT
                if para.alignment == 1:
                    alignment = TA_CENTER
                elif para.alignment == 2:
                    alignment = TA_RIGHT
                elif para.alignment == 3:
                    alignment = TA_JUSTIFY
                
                # Get font size and style
                style = ParagraphStyle(
                    'CustomStyle',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=14,
                    alignment=alignment,
                )
                
                # Check for bold and italic
                runs = para.runs
                text = ''
                for run in runs:
                    if run.bold:
                        text += f'<b>{run.text}</b>'
                    elif run.italic:
                        text += f'<i>{run.text}</i>'
                    else:
                        text += run.text
                
                if not text:
                    text = para.text
                
                story.append(Paragraph(text, style))
                story.append(Spacer(1, 0.2*inch))
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                story.append(Paragraph(' | '.join(row_text), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        pdf_doc.build(story)
        return True
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_docx_to_pdf.py input.docx output.pdf")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"❌ Input file does not exist: {input_file}")
        sys.exit(1)
    
    if convert_docx_to_pdf(input_file, output_file):
        print(f"✅ Successfully converted: {output_file}")
        sys.exit(0)
    else:
        print("❌ Conversion failed")
        sys.exit(1)