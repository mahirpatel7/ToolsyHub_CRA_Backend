# # convert_pdf_to_docx.py
# import sys
# from pdf2docx import Converter
# import os

# def convert(pdf_file, docx_file):
#     cv = Converter(pdf_file)
#     cv.convert(docx_file, start=0, end=None)
#     cv.close()

# if __name__ == "__main__":
#     if len(sys.argv) != 3:
#         print("Usage: python convert_pdf_to_docx.py input.pdf output.docx")
#         sys.exit(1)

#     input_pdf = sys.argv[1]
#     output_docx = sys.argv[2]

#     if not os.path.exists(input_pdf):
#         print("❌ Input PDF does not exist:", input_pdf)
#         sys.exit(1)

#     convert(input_pdf, output_docx)




# convert_pdf_to_docx.py
import sys
import os
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_pdf_to_word(pdf_path, docx_path):
    """
    Convert PDF to DOCX using PyPDF2 + python-docx
    - No system dependencies
    - Fast processing
    - Works on Render.com
    """
    try:
        print(f"📄 Extracting text from PDF: {pdf_path}")
        
        # Extract text from PDF
        pdf_reader = PdfReader(pdf_path)
        
        # Create new Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Converted from PDF', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Extract and add content from each page
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            
            if text.strip():
                # Add page break between pages (except first)
                if page_num > 1:
                    doc.add_page_break()
                
                # Add page number
                page_heading = doc.add_heading(f'Page {page_num}', level=2)
                
                # Split text into paragraphs and add to document
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text)
                        para.paragraph_format.line_spacing = 1.15
        
        # Save the document
        doc.save(docx_path)
        
        # Verify file was created
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
            print(f"✅ Successfully converted to: {docx_path}")
            return True
        else:
            print("❌ DOCX file was not created properly")
            return False
            
    except Exception as e:
        print(f"❌ Conversion Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_pdf_to_docx.py input.pdf output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"❌ Input PDF does not exist: {input_file}")
        sys.exit(1)
    
    # Convert
    if convert_pdf_to_word(input_file, output_file):
        print(f"✅ Conversion successful!")
        sys.exit(0)
    else:
        print("❌ Conversion failed")
        sys.exit(1)